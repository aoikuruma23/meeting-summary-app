import os
import json
from datetime import datetime, timezone, timedelta
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models.meeting import Meeting

class ExportService:
    def __init__(self):
        self.export_dir = os.path.join("uploads", "exports")
        print(f"DEBUG: エクスポートディレクトリ作成開始: {self.export_dir}")
        
        # 親ディレクトリ（uploads）の存在確認
        uploads_dir = "uploads"
        if not os.path.exists(uploads_dir):
            print(f"DEBUG: uploadsディレクトリが存在しません。作成します: {uploads_dir}")
            os.makedirs(uploads_dir, exist_ok=True)
        
        # エクスポートディレクトリの作成
        try:
            os.makedirs(self.export_dir, exist_ok=True)
            print(f"DEBUG: エクスポートディレクトリ作成完了: {self.export_dir}")
        except Exception as e:
            print(f"DEBUG: エクスポートディレクトリ作成エラー: {e}")
            raise
        
        # ディレクトリの存在確認
        print(f"DEBUG: エクスポートディレクトリ存在確認: {os.path.exists(self.export_dir)}")
        
        # ディレクトリの権限確認
        try:
            print(f"DEBUG: エクスポートディレクトリ権限確認: {os.access(self.export_dir, os.W_OK)}")
            print(f"DEBUG: エクスポートディレクトリ読み取り権限: {os.access(self.export_dir, os.R_OK)}")
        except Exception as e:
            print(f"DEBUG: 権限確認エラー: {e}")
        
        # ディレクトリの内容確認
        try:
            files = os.listdir(self.export_dir)
            print(f"DEBUG: エクスポートディレクトリの内容: {files}")
        except Exception as e:
            print(f"DEBUG: ディレクトリ内容確認エラー: {e}")
        
        # 日本語フォントの設定
        self._setup_japanese_fonts()
    
    def _setup_japanese_fonts(self):
        """日本語フォントを設定"""
        try:
            # 利用可能なフォントを確認
            print(f"利用可能なフォント: {pdfmetrics.getRegisteredFontNames()}")
            
            # 1) まずはCIDフォント（組み込み）を登録して即時利用可能にする
            #    これらは外部ファイル不要で、日本語の表示に十分
            try:
                pdfmetrics.registerFont(UnicodeCIDFont('HeiseiKakuGo-W5'))
                self.japanese_font_name = 'HeiseiKakuGo-W5'
                print("日本語CIDフォントを使用: HeiseiKakuGo-W5")
                return
            except Exception as e:
                print(f"CIDフォント登録失敗(HeiseiKakuGo-W5): {e}")
            try:
                pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))
                self.japanese_font_name = 'HeiseiMin-W3'
                print("日本語CIDフォントを使用: HeiseiMin-W3")
                return
            except Exception as e:
                print(f"CIDフォント登録失敗(HeiseiMin-W3): {e}")
            
            # システムフォントの詳細調査
            print("=== システムフォント詳細調査 ===")
            try:
                import subprocess
                # fc-listコマンドで日本語フォントを検索
                result = subprocess.run(['fc-list', ':lang=ja'], capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    print(f"fc-list日本語フォント結果: {result.stdout[:500]}...")
                else:
                    print(f"fc-listエラー: {result.stderr}")
            except Exception as e:
                print(f"fc-list実行エラー: {e}")
            
            # フォントディレクトリの存在確認
            font_dirs = [
                "/usr/share/fonts",
                "/usr/local/share/fonts",
                "/System/Library/Fonts",
                "/Library/Fonts"
            ]
            
            for font_dir in font_dirs:
                if os.path.exists(font_dir):
                    print(f"フォントディレクトリ存在: {font_dir}")
                    try:
                        files = os.listdir(font_dir)
                        japanese_files = [f for f in files if any(char in f.lower() for char in ['japanese', 'japan', 'jp', 'gothic', 'mincho', 'ヒラギノ', 'メイリオ'])]
                        if japanese_files:
                            print(f"日本語フォントファイル発見 ({font_dir}): {japanese_files[:5]}")
                    except Exception as e:
                        print(f"ディレクトリ読み取りエラー ({font_dir}): {e}")
                else:
                    print(f"フォントディレクトリ不存在: {font_dir}")
            
            # 日本語フォントの優先順位
            japanese_fonts = [
                'IPAexGothic',
                'IPAexMincho',
                'TakaoGothic',
                'VL-Gothic-Regular',
                'NotoSansCJK-Regular',
                'NotoSansCJKjp-Regular',
                'NotoSansJP-Regular',
                'SourceHanSans-Regular'
            ]
            
            # 利用可能な日本語フォントを探す
            available_font = None
            for font_name in japanese_fonts:
                try:
                    if font_name in pdfmetrics.getRegisteredFontNames():
                        available_font = font_name
                        print(f"日本語フォントを使用: {font_name}")
                        break
                except:
                    continue
            
            if not available_font:
                # 日本語フォントが見つからない場合、フォントファイルを直接埋め込む
                print("日本語フォントが見つからないため、フォントファイルを埋め込みます")
                try:
                    # 複数の日本語フォントファイルを試行
                    font_files = [
                        os.path.join("fonts", "BIZ-UDGothicR.ttc"),
                        "/usr/share/fonts/truetype/ipafont-gothic/ipag.ttf",
                        "/usr/share/fonts/truetype/ipafont-mincho/ipam.ttf",
                        "/usr/share/fonts/opentype/ipafont-gothic/IPAGothic.ttf",
                        "/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc",  # macOS
                        "/Library/Fonts/Arial Unicode MS.ttf",  # macOS
                        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                        "/usr/share/fonts/opentype/noto/NotoSansCJKjp-Regular.otf",
                        "/usr/share/fonts/truetype/noto/NotoSansJP-Regular.ttf",
                        os.path.join("fonts", "NotoSansCJK-Regular.ttc"),
                        os.path.join("fonts", "NotoSansJP-Regular.otf")
                    ]
                    
                    font_registered = False
                    for font_path in font_files:
                        try:
                            # 絶対パスと相対パスの両方を試行
                            paths_to_try = [font_path]
                            if not os.path.isabs(font_path):
                                # 相対パスの場合、現在の作業ディレクトリからの絶対パスも試行
                                current_dir = os.getcwd()
                                absolute_path = os.path.join(current_dir, font_path)
                                paths_to_try.append(absolute_path)
                                # backendディレクトリからのパスも試行
                                backend_path = os.path.join(current_dir, "backend", font_path)
                                paths_to_try.append(backend_path)
                            
                            for try_path in paths_to_try:
                                if os.path.exists(try_path):
                                    # フォントファイルを登録
                                    font_name = f"JapaneseFont_{os.path.basename(try_path)}"
                                    pdfmetrics.registerFont(TTFont(font_name, try_path))
                                    available_font = font_name
                                    font_registered = True
                                    print(f"日本語フォントファイルを埋め込みました: {try_path}")
                                    break
                                else:
                                    print(f"フォントファイルが存在しません: {try_path}")
                            
                            if font_registered:
                                break
                                
                        except Exception as font_error:
                            print(f"フォントファイル埋め込みエラー ({font_path}): {font_error}")
                            continue
                    
                    if not font_registered:
                        # システムフォントディレクトリを再帰的に検索
                        print("システムフォントディレクトリを検索中...")
                        system_font_dirs = [
                            "/usr/share/fonts",
                            "/usr/local/share/fonts",
                            "/System/Library/Fonts",
                            "/Library/Fonts",
                            "/usr/X11/lib/X11/fonts"
                        ]
                        
                        for font_dir in system_font_dirs:
                            if os.path.exists(font_dir):
                                try:
                                    for root, dirs, files in os.walk(font_dir):
                                        for file in files:
                                            if file.lower().endswith(('.ttf', '.otf', '.ttc')):
                                                font_path = os.path.join(root, file)
                                                try:
                                                    # 日本語フォントかどうかを判定（ファイル名に日本語が含まれているか）
                                                    if any(char in file for char in ['japanese', 'japan', 'jp', 'gothic', 'mincho', 'ヒラギノ', 'メイリオ']):
                                                        font_name = f"SystemFont_{os.path.basename(font_path)}"
                                                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                                                        available_font = font_name
                                                        font_registered = True
                                                        print(f"システムフォントを発見・登録しました: {font_path}")
                                                        break
                                                except Exception as font_error:
                                                    print(f"フォント登録エラー ({font_path}): {font_error}")
                                                    continue
                                        if font_registered:
                                            break
                                    if font_registered:
                                        break
                                except Exception as walk_error:
                                    print(f"フォントディレクトリ検索エラー ({font_dir}): {walk_error}")
                                    continue
                    
                    if not font_registered:
                        # 最後の手段: CIDフォントをもう一度試す
                        try:
                            pdfmetrics.registerFont(UnicodeCIDFont('HeiseiKakuGo-W5'))
                            available_font = 'HeiseiKakuGo-W5'
                            font_registered = True
                            print("フォールバックでCIDフォントを使用: HeiseiKakuGo-W5")
                        except Exception:
                            pass
                        if not font_registered:
                            try:
                                pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))
                                available_font = 'HeiseiMin-W3'
                                font_registered = True
                                print("フォールバックでCIDフォントを使用: HeiseiMin-W3")
                            except Exception:
                                pass
                    
                    if not font_registered:
                        raise Exception("日本語フォントが利用できません。PDFの日本語表示ができません。")
                        
                except Exception as embed_error:
                    print(f"フォント埋め込み処理エラー: {embed_error}")
                    # エラーを再発生させて、PDF生成を中止
                    raise embed_error
            
            if available_font:
                # スタイルで日本語フォントを使用
                self.japanese_font_name = available_font
                print(f"日本語フォント設定完了: {available_font}")
            else:
                # 念のためのフォールバック
                self.japanese_font_name = 'HeiseiKakuGo-W5'
                try:
                    pdfmetrics.registerFont(UnicodeCIDFont(self.japanese_font_name))
                    print(f"CIDフォントにフォールバック: {self.japanese_font_name}")
                except Exception:
                    raise Exception("日本語フォントが利用できません。PDFの日本語表示ができません。")
                
        except Exception as e:
            print(f"日本語フォント設定エラー: {e}")
            print("エラー: 日本語フォントの設定に失敗しました。PDFの日本語表示ができません。")
            # エラーを確実に発生させて、PDF生成を中止
            raise Exception(f"日本語フォントの設定に失敗しました: {str(e)}")
        
        # フォント設定の最終確認
        print(f"最終的なフォント設定: {self.japanese_font_name}")
        print(f"利用可能なフォント一覧: {pdfmetrics.getRegisteredFontNames()}")
        print("=== フォント設定完了 ===")
    
    def jst_now(self):
        """日本時間（JST）の現在時刻を返す"""
        jst = timezone(timedelta(hours=9))
        return datetime.now(jst)
    
    def format_jst_datetime(self, dt):
        """日本時間でフォーマット"""
        if dt:
            # データベースの時間がUTCの場合、JSTに変換
            if dt.tzinfo is None:
                # タイムゾーン情報がない場合はUTCとして扱い、JSTに変換
                utc_time = dt.replace(tzinfo=timezone.utc)
                jst_time = utc_time.astimezone(timezone(timedelta(hours=9)))
            else:
                # 既にタイムゾーン情報がある場合はJSTに変換
                jst_time = dt.astimezone(timezone(timedelta(hours=9)))
            
            return jst_time.strftime('%Y年%m月%d日 %H:%M')
        else:
            # 現在時刻を使用
            return self.jst_now().strftime('%Y年%m月%d日 %H:%M')
    
    def safe_text(self, text):
        """テキストを安全に処理"""
        if not text:
            return ""
        
        # 改行文字を削除
        text = text.replace('\n', ' ')
        # 特殊文字をエスケープ
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        return text
    
    async def export_to_pdf(self, meeting) -> str:
        """要約をPDF形式でエクスポート"""
        try:
            print(f"DEBUG: PDFエクスポート開始 - meeting_id: {meeting.id}")
            
            # 日本語フォントの再確認
            if self.japanese_font_name == 'Helvetica':
                print("エラー: 日本語フォントが利用できません。PDFの日本語表示ができません。")
                raise Exception("日本語フォントが利用できません。PDFの日本語表示ができません。")
            
            print(f"DEBUG: 使用フォント確認: {self.japanese_font_name}")
            
            # ファイル名を生成
            timestamp = self.jst_now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_{meeting.id}_{timestamp}.pdf"
            file_path = os.path.join(self.export_dir, filename)
            
            print(f"DEBUG: PDFファイルパス: {file_path}")
            print(f"DEBUG: エクスポートディレクトリ: {self.export_dir}")
            
            # PDFドキュメントを作成
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            story = []
            
            # スタイルを定義（日本語フォント対応）
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1,  # 中央揃え
                fontName=self.japanese_font_name
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                fontName=self.japanese_font_name
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6,
                fontName=self.japanese_font_name
            )
            
            # タイトルを追加
            story.append(Paragraph(self.safe_text("議事録要約"), title_style))
            story.append(Spacer(1, 20))
            
            # 会議情報を追加
            story.append(Paragraph(self.safe_text(f"会議タイトル: {meeting.title}"), heading_style))
            story.append(Paragraph(self.safe_text(f"作成日時: {self.format_jst_datetime(meeting.created_at)}"), normal_style))
            story.append(Paragraph(self.safe_text(f"議事録ID: {meeting.id}"), normal_style))
            story.append(Spacer(1, 20))
            
            # 要約内容を追加
            story.append(Paragraph(self.safe_text("要約内容"), heading_style))
            
            # 要約テキストを段落に分割して追加
            if meeting.summary:
                summary_paragraphs = meeting.summary.split('\n')
                for paragraph in summary_paragraphs:
                    if paragraph.strip():
                        story.append(Paragraph(self.safe_text(paragraph), normal_style))
                        story.append(Spacer(1, 6))
            
            # PDFを生成
            doc.build(story)
            
            print(f"DEBUG: PDF生成完了 - ファイルパス: {file_path}")
            
            # ファイル生成後の詳細確認
            if os.path.exists(file_path):
                file_size = os.path.getsize(file_path)
                print(f"DEBUG: PDFファイル存在確認: 成功")
                print(f"DEBUG: PDFファイルサイズ: {file_size} bytes")
                
                # ファイルの権限確認
                try:
                    read_ok = os.access(file_path, os.R_OK)
                    write_ok = os.access(file_path, os.W_OK)
                    print(f"DEBUG: PDFファイル読み取り権限: {read_ok}")
                    print(f"DEBUG: PDFファイル書き込み権限: {write_ok}")
                except Exception as e:
                    print(f"DEBUG: ファイル権限確認エラー: {e}")
                
                # ファイルの内容確認（最初の数バイト）
                try:
                    with open(file_path, 'rb') as f:
                        header = f.read(10)
                        print(f"DEBUG: PDFファイルヘッダー: {header}")
                except Exception as e:
                    print(f"DEBUG: ファイル内容確認エラー: {e}")
            else:
                print(f"DEBUG: PDFファイル存在確認: 失敗 - ファイルが存在しません")
                raise Exception("PDFファイルの生成に失敗しました")
            
            return file_path
            
        except Exception as e:
            print(f"DEBUG: PDFエクスポートエラー: {str(e)}")
            import traceback
            print(f"DEBUG: スタックトレース: {traceback.format_exc()}")
            raise
    
    async def export_to_docx(self, meeting) -> str:
        """要約をWord形式でエクスポート"""
        try:
            # ファイル名を生成
            timestamp = self.jst_now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_{meeting.id}_{timestamp}.docx"
            file_path = os.path.join(self.export_dir, filename)
            
            # Wordドキュメントを作成
            doc = Document()
            
            # タイトルを追加
            title = doc.add_heading("議事録要約", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 会議情報を追加
            doc.add_heading("会議情報", level=1)
            doc.add_paragraph(f"会議タイトル: {meeting.title}")
            doc.add_paragraph(f"作成日時: {self.format_jst_datetime(meeting.created_at)}")
            doc.add_paragraph(f"議事録ID: {meeting.id}")
            
            # 要約内容を追加
            doc.add_heading("要約内容", level=1)
            
            # 要約を段落に分割して追加
            paragraphs = meeting.summary.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Wordファイルを保存
            doc.save(file_path)
            
            return file_path
            
        except Exception as e:
            print(f"ERROR: Wordエクスポート失敗: {str(e)}")
            raise Exception(f"Wordエクスポートに失敗しました: {str(e)}")
    
    def cleanup_old_exports(self, max_age_hours: int = 24):
        """古いエクスポートファイルを削除"""
        try:
            current_time = datetime.now()
            for filename in os.listdir(self.export_dir):
                file_path = os.path.join(self.export_dir, filename)
                if os.path.isfile(file_path):
                    file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                    if (current_time - file_time).total_seconds() > max_age_hours * 3600:
                        os.remove(file_path)
                        print(f"古いエクスポートファイルを削除: {filename}")
        except Exception as e:
            print(f"エクスポートファイルのクリーンアップに失敗: {str(e)}") 